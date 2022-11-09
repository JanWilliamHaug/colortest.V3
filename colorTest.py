# Instructions on how to use the program:
# 1. run program
# 2. click on button "Choose document" and choose a document
# 3. click on button "Generate Report" (this will print it out on the terminal and add it
# to a word document)
# 4. click on button "Choose document" if you want to add more documents, and click on the
#   button "Generate Report" again
# 5. When you don't wanna add anymore documents, click End Program
# 6. Contents found by the program should all now have been outputted in the terminal and
# in the word document generated by the program

# if you use python version 3.9 no changes in code is needed,
# if you use Version 3.7: You have to change "tuple" to "Tuple" on line

import docx
from docx import Document
from docx.shared import RGBColor
from tkinter import *
from tkinter import filedialog

from typing import Tuple


def readtxt(filename, color: Tuple[int, int, int]):
    doc = docx.Document(filename)
    text10 = ""
    fullText = []
    new = []
    for para in doc.paragraphs:

        # Getting the colored words from the doc
        if (getcoloredTxt(para.runs, color)):

            # Concatenating list of runs between the colored text to single a string

            sentence = "".join(r.text for r in para.runs)
            fullText.append(sentence)
            print(sentence)
            text10 = sentence
            #new = (sentence.replace (']', ']\n\n'))
            parent.append("".join(r.text for r in para.runs))



    #print(fullText)
    global filtered_L
    global hasChild
    global fullText2
    filtered_L = [value for value in fullText if "[" not in value]
    hasChild = [value for value in fullText if "[" in value]
    fullText2 = [value for value in fullText]


    #print(filtered_L)
    #print(fullText)
    return fullText, filtered_L, hasChild

def getcoloredTxt(runs, color):

    coloredWords, word = [], ""
    for run in runs:
        if run.font.color.rgb == RGBColor(*color):
            word += str(run.text)


        elif word != "":
            coloredWords.append(word)
            sentences.append(word)
            parents.append(word)
            word = ""

    #if word == "":
    #   noChild.append(word)

    if word != "":
        coloredWords.append(word + "\n")
        #word = removeAfter(word)
        child.append(word)
        withChild.append(word)



    return coloredWords

def openFile():
    global filepath
    global filepath2
    filepath = filedialog.askopenfilename(initialdir="/",
                                          title="",
                                          filetypes= (("word documents","*.docx"),
                                                      ("all files","*.*")))
    file = open(filepath,'r')
    #print(filepath)
    file.close()
    filepath2 = str(filepath)
    #filepath2 = '"' + filepath + '"'
    print(filepath2)

    return filepath2

def generateReport():
    fullText = readtxt(filename=filepath2,
                       color=(255, 0, 0))
    #filtered_L = readtxt(filename=filepath2,
    #                   color=(255, 0, 0))
    fullText10 = str(fullText)
    s = ''.join(fullText10)
    w = (s.replace (']', ']\n\n'))
    #w = (w.replace ('\n[', '['))
    #print('\n' + w)
    paragraph = report3.add_paragraph()
    runner = paragraph.add_run("\n" + filepath2 + "\n")
    runner.bold = True #makes the header bold
    w = (w.replace ('([', ''))
    w = (w.replace (',', ''))
    w = (w.replace ('' '', ''))

    #print(w)
    #print(fullText)


    table = report3.add_table(rows=1, cols=2)

    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Parent Tag'
    row[1].text = 'Child Tag/Tags'



    # Adding style to a table
    table.style = 'Colorful List'

    # Now save the document to a location
    report3.save('report3.docx')

    #print(filtered_L)
    #print(fullText)
    #print(fullText2)
    e = 0
    #print(child)
    child2 = removeAfter(child) #removes everything after the child tag if there is anything to remove
    #print(child2)
    while sentences:
        row = table.add_row().cells # Adding a row and then adding data in it.
        row[0].text = sentences[0]
        #green = paragraph.add_run(sentences[0] + "    ")
        #paragraph.add_run("\n\n")
        #green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
        #green.bold = True
        sentences.remove(sentences[0])
        #print(fullText)
        #print(filtered_L)


        if e < len(fullText2):

            if fullText2[e] in filtered_L:

                #row = table.add_row().cells # Adding a row and then adding data in it.
                # print("no child")
                #print("yes")

                row[1].text = " "

                e += 1

            elif fullText2[e] not in filtered_L:
                #print("has a child")
                #print(fullText[e])
                #row[1].text = "Has no child tag"
                if child2:
                    #red = paragraph.add_run(child[0])
                    row[1].text = child2[0]
                    #paragraph.add_run("\n\n")
                    #red.bold = True
                    #red.font.color.rgb = RGBColor(255, 0, 0)
                    child2.remove(child2[0])

                    e += 1



    while sentences:
        row = table.add_row().cells # Adding a row and then adding data in it.
        row[0].text = sentences[0]

    while child2:
        row = table.add_row().cells # Adding a row and then adding data in it.
        row[0].text = sentences[0]

    #green = paragraph.add_run(sentences[0])
    #paragraph.add_run("\n\n")
    #green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
    #green.bold = True
    #sentences.remove(sentences[0])
    child2.clear()
    sentences.clear()
    child.clear()
    #paragraph.add_run(f)
    report3.save('report3.docx')


def removeAfter(childtags): #removes everything after the child tag, example "pass"
    seperator = ']'
    childAfter = [i.rsplit(']', 1)[0] + seperator for i in childtags]

    return childAfter


if __name__ == '__main__':
    # Creates a word document, saves it as "report 3, and also adds a heading
    report3 = Document()
    report3.add_heading('Report', 0) #create word document
    paragraph = report3.add_paragraph()
    report3.save('report3.docx')

    # declaring different lists that will be used to store, tags and sentences
    sentences = []
    parent = []
    child = []
    noChild = []
    withChild = []
    parents = []

    #fulltext9 = ''.join(parent)
    #g = str(fulltext9)
    #f = (g.replace (']', ']\n\n'))

    # Creates the gui
    window = Tk(className=' TARGEST')

    # set window size
    window.geometry("150x100")
    # Creates button 1
    button = Button(text="Choose Document",command=openFile)
    button.pack()
    # Creates button 2
    Button(window, text="Generate Report ", command=generateReport).pack()
    # Creates button 3
    button = Button(text="End Program",command=window.destroy)
    button.pack()

    window.mainloop()

    #print(sentences) #parent tags
    #print(child) #child tags
    #print(parent)
    #print(noChild)
    #print(withChild)
    #print(parents)



    #print(fullText)
    #lister2 = [fullText]
    #d = ''.join(fullText)
    #print(fullText)
    #print(d)
    #words = d.split('')
    #words2 = d.split('')
    #print(words)




